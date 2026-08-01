from io import BytesIO
from datetime import date

from docx import Document


def build_evidence_pack(intake: dict, result: dict) -> BytesIO:
    """
    Builds a Word document from the intake + AI Compass result and
    returns it as an in-memory BytesIO buffer (no file written to disk).
    """
    doc = Document()

    # --- Header / use case summary ---
    doc.add_heading("AI Compass — Evidence Pack", level=0)
    doc.add_paragraph(f"Date generated: {date.today().isoformat()}")

    doc.add_heading("Use case summary", level=1)
    doc.add_paragraph(f"Business problem: {intake.get('business_problem', '')}")
    doc.add_paragraph(f"Proposed AI tool/model: {intake.get('proposed_tool', '')}")
    doc.add_paragraph(f"Data sensitivity: {intake.get('data_sensitivity', '')}")
    doc.add_paragraph(f"Users affected: {intake.get('users_affected', '')}")
    doc.add_paragraph(f"Expected value: {intake.get('expected_value', '')}")
    doc.add_paragraph(f"Estimated usage: {intake.get('estimated_usage', '')}")
    impact = ", ".join(intake.get("decision_impact", [])) or "None specified"
    doc.add_paragraph(f"Decision affects: {impact}")

    # --- Governance decision ---
    governance = result.get("governance_decision", {})
    doc.add_heading("Governance decision", level=1)
    decision_label = governance.get("decision", "").replace("_", " ").title()
    doc.add_paragraph(f"Decision: {decision_label}")
    doc.add_paragraph(governance.get("rationale", ""))

    if governance.get("conditions"):
        doc.add_heading("Conditions", level=2)
        for condition in governance["conditions"]:
            doc.add_paragraph(condition, style="List Bullet")

    if governance.get("suggested_next_review"):
        doc.add_paragraph(f"Suggested next review: {governance['suggested_next_review']}")

    # --- Risk assessment table ---
    risks = result.get("risk_assessment", {})
    doc.add_heading("Risk assessment", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Category"
    header_cells[1].text = "Level"
    header_cells[2].text = "Rationale"

    for name, risk in risks.items():
        row_cells = table.add_row().cells
        row_cells[0].text = name.replace("_", " ").title()
        row_cells[1].text = risk.get("level", "").title()
        row_cells[2].text = risk.get("rationale", "")

    # --- Evaluation plan ---
    plan = result.get("evaluation_plan", {})
    doc.add_heading("Evaluation plan", level=1)

    plan_sections = [
        ("Test cases", "test_cases"),
        ("Success metrics", "success_metrics"),
        ("Acceptance thresholds", "acceptance_thresholds"),
        ("Red-team scenarios", "red_team_scenarios"),
        ("UAT checklist", "uat_checklist"),
    ]
    for label, key in plan_sections:
        doc.add_heading(label, level=2)
        for item in plan.get(key, []):
            doc.add_paragraph(item, style="List Bullet")

    # --- Save to memory, not disk ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
